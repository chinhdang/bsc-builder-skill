"""
BSC DOCX Export — Helper functions & low-level builders.
Imported by docx-export-script.py. Do not run directly.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime
import os

# ── Color Palette ──────────────────────────────────────────────────
DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
MED_BLUE  = RGBColor(0x2E, 0x75, 0xB6)
LIGHT_BLUE_HEX = "DEEAF1"
DARK_BLUE_HEX  = "1F4E79"
MED_BLUE_HEX   = "2E75B6"
GRAY       = RGBColor(0x59, 0x59, 0x59)
DARK_GRAY  = RGBColor(0x44, 0x44, 0x44)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x00, 0x00, 0x00)
LIGHT_GRAY_HEX = "F2F2F2"
BORDER_GRAY    = "BFBFBF"
GREEN_HEX = "E2EFDA"
RED_HEX   = "FCE4EC"

# Semantic colors for risk/status indicators
COLORS = {
    "red":          "C0392B",
    "red_light":    "FADBD8",
    "orange":       "CA6F1E",
    "orange_light": "FDEBD0",
    "gold":         "B7950B",
    "gold_light":   "FEF9E7",
    "green":        "1E8449",
    "green_light":  "D5F5E3",
}


# ── Cell Helpers ───────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top, bottom, left, right with (size, color)."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, (sz, color) in kwargs.items():
        el = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="{sz}" '
            f'w:space="0" w:color="{color}"/>'
        )
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_vertical_alignment(cell, align="center"):
    """Set vertical alignment: top, center, bottom."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{align}"/>')
    tcPr.append(vAlign)


def set_cell_margins(cell, top=0, bottom=0, left=100, right=100):
    """Set cell margins in twips."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:start w:w="{left}" w:type="dxa"/>'
        f'<w:end w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(margins)


def set_cell_text_multiline(cell, text, size=10, color=None, bold=False, margins=True):
    """Write text into a cell, splitting on \\n into separate paragraphs."""
    if color is None:
        color = BLACK
    lines = str(text).split("\n")
    # Use the first existing paragraph
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        add_run(p, line, size=size, color=color, bold=bold)
    if margins:
        set_cell_margins(cell, top=40, bottom=40, left=80, right=80)


# ── Paragraph Helpers ──────────────────────────────────────────────

def add_run(paragraph, text, size=10, bold=False, italic=False,
            color=None, font_name=None):
    """Add a styled run to paragraph."""
    if color is None:
        color = BLACK
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    if font_name:
        run.font.name = font_name
    return run


def add_mixed_run(doc, runs_spec, alignment=None, space_before=0, space_after=6):
    """Add a paragraph with multiple runs of different styles.

    runs_spec: list of dicts, each with keys:
        text (required), size, bold, italic, color (RGBColor)
    Example:
        add_mixed_run(doc, [
            {"text": "Nhận định: ", "bold": True},
            {"text": "Thị trường đang ở giai đoạn switching...", "italic": True, "color": GRAY},
        ])
    """
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    for spec in runs_spec:
        add_run(
            p,
            spec["text"],
            size=spec.get("size", 10),
            bold=spec.get("bold", False),
            italic=spec.get("italic", False),
            color=spec.get("color", BLACK),
        )
    return p


def add_styled_paragraph(doc, text, size=10, bold=False, italic=False,
                          color=None, alignment=None, space_before=0,
                          space_after=6):
    """Add a paragraph with consistent styling."""
    if color is None:
        color = BLACK
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    add_run(p, text, size=size, bold=bold, italic=italic, color=color)
    return p


def add_paragraph_with_bottom_border(doc, text, size=13, bold=True,
                                     color=None, border_color="2E75B6"):
    """H2-style paragraph with a bottom border line (light blue)."""
    if color is None:
        color = MED_BLUE
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(7)
    add_run(p, text, size=size, bold=bold, color=color)

    # Apply bottom border via pPr/pBdr XML
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="{border_color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


# ── Section / Layout Helpers ───────────────────────────────────────

def add_section_banner(doc, text):
    """Full-width dark-blue banner for H1-level section headings."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Expand to full page width
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="0" w:type="auto"/>')
    tblPr.append(tblW)

    cell = table.rows[0].cells[0]
    set_cell_shading(cell, DARK_BLUE_HEX)
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, text, size=14, bold=True, color=WHITE)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_section_intro(doc, text):
    """Italic gray intro paragraph placed below a section banner."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    add_run(p, text, size=10, italic=True, color=GRAY)
    return p


def add_insight_box(doc, text):
    """Left-border insight box (primary blue left border + light blue bg)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, LIGHT_BLUE_HEX)
    set_cell_margins(cell, top=80, bottom=80, left=150, right=150)

    # Thick blue left border only
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{MED_BLUE_HEX}"/>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

    p = cell.paragraphs[0]
    add_run(p, text, size=10, italic=True, color=DARK_GRAY)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


# ── Header / Footer ───────────────────────────────────────────────

def _add_border_paragraph(paragraph, border_type="bottom", color="2E75B6", sz=4):
    """Add top or bottom border to a paragraph via XML."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:{border_type} w:val="single" w:sz="{sz}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def setup_header_footer(section, doc_title, company_name, date_str):
    """Add professional header and footer to a document section.

    Header: doc_title (left, bold, blue) | company_name (right, gray)
            with bottom border.
    Footer: "Tài liệu nội bộ" (left) | page X/Y (center) | date (right)
            with top border.
    """
    # ── Header ──
    header = section.header
    header.is_linked_to_previous = False
    # Clear default paragraph
    for p in header.paragraphs:
        p.clear()
    h_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

    # Use a 1-row, 2-col table inside header for left/right alignment
    h_table = header.add_table(rows=1, cols=2, width=Cm(17))
    h_table.autofit = True
    left_cell = h_table.rows[0].cells[0]
    right_cell = h_table.rows[0].cells[1]

    # Left: title
    lp = left_cell.paragraphs[0]
    add_run(lp, doc_title, size=9, bold=True, color=DARK_BLUE)

    # Right: company
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(rp, company_name, size=9, color=GRAY)

    # Bottom border on the table row cells
    for cell in [left_cell, right_cell]:
        set_cell_border(cell, bottom=("4", MED_BLUE_HEX))
        set_cell_margins(cell, top=40, bottom=40, left=0, right=0)

    # Remove the blank paragraph that was there before
    if h_para != header.paragraphs[-1]:
        p_elem = h_para._element
        p_elem.getparent().remove(p_elem)

    # ── Footer ──
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()

    f_table = footer.add_table(rows=1, cols=3, width=Cm(17))
    f_table.autofit = True
    f_cells = f_table.rows[0].cells

    # Left: "Tài liệu nội bộ"
    lp2 = f_cells[0].paragraphs[0]
    add_run(lp2, "Tài liệu nội bộ", size=8, color=GRAY)

    # Center: page X / Y via field code
    cp = f_cells[1].paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cp, "Trang ", size=8, color=GRAY)
    # PAGE field
    fld_page = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
    )
    r1 = cp.add_run()
    r1._element.append(fld_page)
    r2 = cp.add_run()
    r2._element.append(parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
    ))
    r3 = cp.add_run()
    r3._element.append(parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
    ))
    add_run(cp, " / ", size=8, color=GRAY)
    # NUMPAGES field
    r4 = cp.add_run()
    r4._element.append(parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
    ))
    r5 = cp.add_run()
    r5._element.append(parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>'
    ))
    r6 = cp.add_run()
    r6._element.append(parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
    ))
    # Set size on field runs
    for r in [r1, r2, r3, r4, r5, r6]:
        r.font.size = Pt(8)
        r.font.color.rgb = GRAY

    # Right: date
    rp2 = f_cells[2].paragraphs[0]
    rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(rp2, date_str, size=8, color=GRAY)

    # Top border on footer cells
    for cell in f_cells:
        set_cell_border(cell, top=("4", BORDER_GRAY))
        set_cell_margins(cell, top=40, bottom=40, left=0, right=0)

    # Remove leftover empty first paragraph
    first_p = footer.paragraphs[0] if footer.paragraphs else None
    last_p  = footer.paragraphs[-1] if footer.paragraphs else None
    if first_p and first_p != last_p and not first_p.text.strip():
        first_p._element.getparent().remove(first_p._element)
