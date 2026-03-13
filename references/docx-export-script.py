"""
BSC Strategic Plan — Professional DOCX Export Script
Generates a polished DOCX with styled tables, callout boxes, cover page,
section banners, header/footer, insight boxes, and strategy cards.

Usage by AI:
  1. Read bsc-data/*.md files
  2. Populate the data dict below with actual content
  3. Run: python3 docx-export-script.py

Dependencies: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime
import os
import sys

# Allow importing helpers from same directory
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx_export_helpers import (
    # Colors
    DARK_BLUE, MED_BLUE, GRAY, DARK_GRAY, WHITE, BLACK,
    LIGHT_BLUE_HEX, DARK_BLUE_HEX, MED_BLUE_HEX,
    LIGHT_GRAY_HEX, BORDER_GRAY, GREEN_HEX, RED_HEX, COLORS,
    # Cell helpers
    set_cell_shading, set_cell_border, set_cell_vertical_alignment,
    set_cell_margins, set_cell_text_multiline,
    # Paragraph helpers
    add_run, add_mixed_run, add_styled_paragraph,
    add_paragraph_with_bottom_border,
    # Section/layout helpers
    add_section_banner, add_section_intro, add_insight_box,
    # Header/footer
    setup_header_footer,
)


# ── Heading Helper ─────────────────────────────────────────────────

def add_heading_styled(doc, text, level=1):
    """Add section heading.

    Level 1 → full-width dark-blue banner (section_banner).
    Level 2 → H2 with bottom border.
    Level 3 → styled paragraph (bold, medium blue).
    """
    if level == 1:
        return add_section_banner(doc, text)
    elif level == 2:
        return add_paragraph_with_bottom_border(doc, text, size=13, bold=True,
                                                color=MED_BLUE)
    else:
        return add_styled_paragraph(doc, text, size=11, bold=True, color=MED_BLUE,
                                    space_before=10, space_after=5)


# ── Component Builders ─────────────────────────────────────────────

def add_cover_page(doc, company_name, subtitle, period, date_str,
                   prepared_by, prepared_for):
    """Professional cover page (no header/footer via different_first_page)."""
    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "BSC STRATEGIC PLAN", size=28, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, company_name, size=20, bold=True, color=MED_BLUE)

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, subtitle, size=14, color=GRAY)

    doc.add_paragraph()

    for text in [
        f"Kỳ chiến lược: {period}",
        f"Chuẩn bị bởi: {prepared_by}",
        f"Dành cho: {prepared_for}",
        f"Ngày: {date_str}",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, text, size=11, color=GRAY)

    doc.add_page_break()


def add_callout_box(doc, title, body_text, bg_color=None):
    """Callout box: single-cell table with colored background."""
    if bg_color is None:
        bg_color = LIGHT_BLUE_HEX
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]

    set_cell_shading(cell, bg_color)
    set_cell_border(cell,
                    top=("1", BORDER_GRAY), bottom=("1", BORDER_GRAY),
                    left=("1", BORDER_GRAY), right=("1", BORDER_GRAY))
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)

    p = cell.paragraphs[0]
    add_run(p, title, size=13, bold=True, color=DARK_BLUE)

    p2 = cell.add_paragraph()
    add_run(p2, body_text, size=10, color=GRAY)

    doc.add_paragraph()
    return table


def add_data_table(doc, headers, rows, col_widths=None):
    """Professional data table with dark-blue header row.

    Cell text supports \\n for multi-line content within a cell.
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, DARK_BLUE_HEX)
        set_cell_margins(cell, top=50, bottom=50, left=80, right=80)
        p = cell.paragraphs[0]
        add_run(p, header, size=10, bold=True, color=WHITE)

    # Data rows
    for r_idx, row in enumerate(rows):
        bg = LIGHT_GRAY_HEX if r_idx % 2 == 1 else None
        for c_idx, cell_text in enumerate(row):
            if c_idx >= len(headers):
                continue
            cell = table.rows[r_idx + 1].cells[c_idx]
            if bg:
                set_cell_shading(cell, bg)
            # Support \n line breaks in cell content
            set_cell_text_multiline(cell, str(cell_text), size=10, color=BLACK)

    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table


def add_comparison_table(doc, title, before_items, after_items):
    """Before/After comparison table with red/green coloring."""
    if title:
        add_heading_styled(doc, title, level=2)

    table = doc.add_table(rows=1 + len(before_items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(["Trước", "Sau"]):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, DARK_BLUE_HEX)
        set_cell_margins(cell, top=50, bottom=50, left=80, right=80)
        add_run(cell.paragraphs[0], h, size=10, bold=True, color=WHITE)

    for r_idx in range(len(before_items)):
        for col_idx, (item, bg) in enumerate(
            [(before_items[r_idx], RED_HEX), (after_items[r_idx], GREEN_HEX)]
        ):
            cell = table.rows[r_idx + 1].cells[col_idx]
            set_cell_shading(cell, bg)
            set_cell_text_multiline(cell, item, size=10, color=BLACK)

    doc.add_paragraph()
    return table


def add_numbered_card(doc, number, title, description, details=None):
    """Feature card: number badge (blue) + description (light blue bg)."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    num_cell = table.rows[0].cells[0]
    num_cell.width = Cm(1.5)
    set_cell_shading(num_cell, MED_BLUE_HEX)
    set_cell_vertical_alignment(num_cell, "center")
    set_cell_margins(num_cell, top=80, bottom=80, left=60, right=60)
    p = num_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, str(number), size=18, bold=True, color=WHITE)

    desc_cell = table.rows[0].cells[1]
    set_cell_shading(desc_cell, LIGHT_BLUE_HEX)
    set_cell_margins(desc_cell, top=80, bottom=80, left=120, right=120)

    add_run(desc_cell.paragraphs[0], title, size=12, bold=True, color=MED_BLUE)

    if description:
        add_run(desc_cell.add_paragraph(), description, size=10, color=DARK_GRAY)

    if details:
        add_run(desc_cell.add_paragraph(), details, size=9.5, color=DARK_GRAY)

    doc.add_paragraph()
    return table


def add_strategy_card(doc, number, title, type_text, rationale):
    """Strategy card: colored header (dark blue, white text) + light blue body.

    Args:
        number: card sequence number
        title:  strategy name/title
        type_text: e.g. "Type: Direct | TOWS: S-O"
        rationale: descriptive body text
    """
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row spans 2 cols via merge
    header_cell = table.rows[0].cells[0].merge(table.rows[0].cells[1])
    set_cell_shading(header_cell, DARK_BLUE_HEX)
    set_cell_margins(header_cell, top=80, bottom=80, left=120, right=120)
    hp = header_cell.paragraphs[0]
    add_run(hp, f"{number}. {title}", size=12, bold=True, color=WHITE)
    if type_text:
        hp2 = header_cell.add_paragraph()
        add_run(hp2, type_text, size=9, color=RGBColor(0xBF, 0xD7, 0xED))

    # Body: number badge (left) + rationale (right)
    num_cell = table.rows[1].cells[0]
    num_cell.width = Cm(1.5)
    set_cell_shading(num_cell, MED_BLUE_HEX)
    set_cell_vertical_alignment(num_cell, "center")
    set_cell_margins(num_cell, top=80, bottom=80, left=60, right=60)
    np_ = num_cell.paragraphs[0]
    np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(np_, str(number), size=18, bold=True, color=WHITE)

    body_cell = table.rows[1].cells[1]
    set_cell_shading(body_cell, LIGHT_BLUE_HEX)
    set_cell_margins(body_cell, top=80, bottom=80, left=120, right=120)
    add_run(body_cell.paragraphs[0], rationale, size=10, color=DARK_GRAY)

    doc.add_paragraph()
    return table


def add_kpi_table(doc, perspective_name, kpis):
    """KPI table for a BSC perspective.

    kpis: list of dicts with keys: name, type, baseline, q1, q2, q3, owner
    """
    add_heading_styled(doc, perspective_name, level=2)
    headers = ["#", "KPI", "Loại", "Baseline", "Q1", "Q2", "Q3", "Owner"]
    rows = []
    for i, kpi in enumerate(kpis, 1):
        rows.append([
            str(i),
            kpi.get("name", ""),
            kpi.get("type", ""),
            kpi.get("baseline", ""),
            kpi.get("q1", ""),
            kpi.get("q2", ""),
            kpi.get("q3", ""),
            kpi.get("owner", ""),
        ])
    add_data_table(doc, headers, rows)


def add_causal_chain_table(doc, chains):
    """Causal chain map table.

    chains: list of dicts with keys: id, foundation, process, customer, financial
    """
    headers = ["Chain", "Foundation (Leading)", "Process", "Customer",
               "Financial (Lagging)"]
    rows = []
    for c in chains:
        rows.append([
            str(c.get("id", "")),
            c.get("foundation", ""),
            c.get("process", ""),
            c.get("customer", ""),
            c.get("financial", ""),
        ])
    add_data_table(doc, headers, rows)


def add_resource_loading(doc, people):
    """Resource loading table.

    people: list of dicts with keys: name, role, allocations, total, status
    """
    for person in people:
        add_heading_styled(doc, f"{person['name']} — {person['role']}", level=3)
        headers = ["Hoạt động", "Phân bổ %"]
        rows = [[a["activity"], f"{a['percent']}%"]
                for a in person["allocations"]]
        rows.append(["TỔNG", f"{person['total']}%"])
        add_data_table(doc, headers, rows)

        status = person.get("status", "Balanced")
        color = MED_BLUE if "Balanced" in status else RGBColor(0xC0, 0x39, 0x2B)
        add_styled_paragraph(doc, f"Status: {status}", size=10, bold=True,
                              color=color)


def add_four_d_table(doc, person_name, current, ideal, quarterly):
    """4D Mix analysis table.

    current/ideal: dict with keys: do, decide, delegate, design
    quarterly: list of dicts with keys: quarter, do, decide, delegate, design
    """
    add_heading_styled(doc, f"4D Mix — {person_name}", level=3)
    headers = ["Quarter", "Do", "Decide", "Delegate", "Design"]
    rows = [["Current", f"{current['do']}%", f"{current['decide']}%",
             f"{current['delegate']}%", f"{current['design']}%"]]
    for q in quarterly:
        rows.append([q["quarter"], f"{q['do']}%", f"{q['decide']}%",
                     f"{q['delegate']}%", f"{q['design']}%"])
    rows.append(["Ideal", f"{ideal['do']}%", f"{ideal['decide']}%",
                 f"{ideal['delegate']}%", f"{ideal['design']}%"])
    add_data_table(doc, headers, rows)


def add_image_with_caption(doc, image_path, caption, width_cm=15):
    """Add image centered with italic caption below."""
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(image_path, width=Cm(width_cm))

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_cap, caption, size=9, italic=True, color=GRAY)
    doc.add_paragraph()


def add_bullet_list(doc, items, size=10, color=None):
    """Add bullet list items."""
    if color is None:
        color = BLACK
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        if p.runs:
            p.runs[0].font.size = Pt(size)
            p.runs[0].font.color.rgb = color
            p.runs[0].text = item
        else:
            add_run(p, item, size=size, color=color)


def add_swot_matrix(doc, strengths, weaknesses, opportunities, threats):
    """2x2 SWOT matrix with semantic colored quadrants."""
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Top header row
    for i, label in enumerate(["Internal", "External"]):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, DARK_BLUE_HEX)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, label if i == 0 else "", size=10, bold=True, color=WHITE)

    # Quadrant config: (row, col, label, bg_hex, text_color)
    quadrants = [
        (1, 0, "STRENGTHS",    COLORS["green_light"],  COLORS["green"]),
        (1, 1, "WEAKNESSES",   COLORS["red_light"],    COLORS["red"]),
        (2, 0, "OPPORTUNITIES", LIGHT_BLUE_HEX,        DARK_BLUE_HEX),
        (2, 1, "THREATS",      COLORS["gold_light"],   COLORS["gold"]),
    ]
    items_map = [strengths, weaknesses, opportunities, threats]

    for idx, (row, col, label, bg, text_color_hex) in enumerate(quadrants):
        cell = table.rows[row].cells[col]
        set_cell_shading(cell, bg)
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)

        label_color = RGBColor(
            int(text_color_hex[0:2], 16),
            int(text_color_hex[2:4], 16),
            int(text_color_hex[4:6], 16),
        )
        add_run(cell.paragraphs[0], label, size=11, bold=True, color=label_color)

        for item in items_map[idx]:
            add_run(cell.add_paragraph(), f"  {item}", size=9.5, color=BLACK)

    doc.add_paragraph()
    return table


# ── Page Setup ─────────────────────────────────────────────────────

def setup_document():
    """Create document with professional page setup and styles."""
    doc = Document()

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        # Cover page gets its own header/footer (blank)
        section.different_first_page_header_footer = True

    for i, (size, color) in enumerate(
        [(16, DARK_BLUE), (13, MED_BLUE), (11, MED_BLUE)], start=1
    ):
        style = doc.styles[f"Heading {i}"]
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(20 if i == 1 else 14)
        style.paragraph_format.space_after = Pt(10 if i == 1 else 7)

    normal = doc.styles["Normal"]
    normal.font.size = Pt(10)
    normal.font.name = "Calibri"
    normal.paragraph_format.space_after = Pt(6)

    return doc


# ── Main Export Function ───────────────────────────────────────────

def export_bsc_plan(data, output_path):
    """Export BSC Strategic Plan to professional DOCX.

    data: dict with all BSC plan content (see DATA_TEMPLATE below)
    output_path: path for output .docx file
    """
    doc = setup_document()
    date_str = data.get("date", datetime.now().strftime("%d/%m/%Y"))
    doc_title = data.get("doc_title", "BSC Strategic Plan")
    company_name = data.get("company_name", "")

    # ─── Cover Page (first page — no header/footer) ───
    add_cover_page(
        doc,
        company_name=company_name,
        subtitle=data.get("subtitle", ""),
        period=data.get("period", "2025-2026"),
        date_str=date_str,
        prepared_by=data.get("prepared_by", "BSC Strategic Planning Consultant"),
        prepared_for=data.get("prepared_for", "Ban Lãnh đạo"),
    )

    # ─── Header / Footer (applied to the main section, skip cover) ───
    section = doc.sections[0]
    setup_header_footer(section, doc_title, company_name, date_str)

    # ─── Executive Summary ───
    add_callout_box(doc, "Tóm tắt", data.get("executive_summary", ""))

    # ─── Summary Dashboard ───
    if "dashboard" in data:
        add_heading_styled(doc, "Summary Dashboard", level=1)
        if "dashboard_intro" in data:
            add_section_intro(doc, data["dashboard_intro"])
        add_data_table(doc, ["Metric", "Value"], data["dashboard"])

    # ─── Company Profile ───
    if "company_profile" in data:
        add_heading_styled(doc, "Company Profile", level=1)
        for line in data["company_profile"]:
            add_styled_paragraph(doc, line, size=10)

    # ─── Vision & Mission ───
    if "vision" in data:
        add_heading_styled(doc, "Vision & Mission", level=1)
        add_styled_paragraph(doc, f"Vision: {data['vision']}", size=10,
                              bold=True, color=DARK_BLUE)
        add_styled_paragraph(doc, f"Mission: {data.get('mission', '')}",
                              size=10, bold=True, color=MED_BLUE)

        if "vision_milestones" in data:
            add_heading_styled(doc, "Vision Milestones", level=2)
            for ms in data["vision_milestones"]:
                add_styled_paragraph(doc, ms, size=10)

    # ─── Products & Services ───
    if "products" in data:
        add_heading_styled(doc, "Products & Services", level=1)
        for i, prod in enumerate(data["products"], 1):
            add_numbered_card(doc, i, prod["name"],
                               prod.get("description", ""),
                               prod.get("details", ""))

    # ─── Reality Check ───
    if "reality_check" in data:
        add_heading_styled(doc, "Reality Check — Key Insights", level=1)
        rc = data["reality_check"]

        if "intro" in rc:
            add_section_intro(doc, rc["intro"])

        if "market_context" in rc:
            add_heading_styled(doc, "Market Context", level=2)
            add_styled_paragraph(doc, rc["market_context"], size=10)

        if "competitive_position" in rc:
            add_heading_styled(doc, "Competitive Position", level=2)
            add_styled_paragraph(doc, rc["competitive_position"], size=10)

        if "critical_risks" in rc:
            add_heading_styled(doc, "Critical Risks", level=2)
            # Severity column uses semantic color via cell text (plain for now)
            add_data_table(doc, ["Risk", "Severity", "Mitigation"],
                           rc["critical_risks"])

        if "perception_gaps" in rc:
            add_heading_styled(doc, "Perception vs Reality", level=2)
            add_data_table(doc, ["Dimension", "User Perception",
                                 "Research Finding", "Gap"],
                           rc["perception_gaps"])

        if "key_insight" in rc:
            add_insight_box(doc, rc["key_insight"])

    # ─── SWOT Matrix ───
    if "swot" in data:
        add_heading_styled(doc, "SWOT Analysis", level=1)
        s = data["swot"]
        add_swot_matrix(doc, s.get("strengths", []), s.get("weaknesses", []),
                        s.get("opportunities", []), s.get("threats", []))

        if "insight" in s:
            add_insight_box(doc, s["insight"])

    # ─── Core Strategies ───
    if "strategies" in data:
        add_heading_styled(doc, "Core Strategies", level=1)
        if "strategies_intro" in data:
            add_section_intro(doc, data["strategies_intro"])
        for i, strat in enumerate(data["strategies"], 1):
            add_strategy_card(
                doc, i, strat["name"],
                f"Type: {strat.get('type', '')} | TOWS: {strat.get('tows', '')}",
                strat.get("rationale", "")
            )

        if "strategies_rejected" in data:
            add_heading_styled(doc, "Strategies Not Selected", level=2)
            for s in data["strategies_rejected"]:
                add_styled_paragraph(doc, f"- {s}", size=10, color=GRAY)

    # ─── Strategic Objectives ───
    if "objectives" in data:
        add_heading_styled(doc, "Strategic Objectives", level=1)
        for perspective in data["objectives"]:
            add_heading_styled(doc, perspective["name"], level=2)
            headers = ["#", "Objective", "KPI", "Q1", "Q2", "Q3"]
            add_data_table(doc, headers, perspective["items"])

            if "causality" in perspective:
                add_insight_box(doc, perspective["causality"])

    # ─── KPIs ───
    if "kpis" in data:
        add_heading_styled(doc, "KPIs", level=1)
        for perspective in data["kpis"]:
            add_kpi_table(doc, perspective["name"], perspective["items"])

    # ─── Causal Chain Map ───
    if "causal_chains" in data:
        add_heading_styled(doc, "Causal Chain Map", level=1)
        add_causal_chain_table(doc, data["causal_chains"])

        if "chain_narrative" in data:
            add_insight_box(doc, data["chain_narrative"])

    # ─── Initiatives ───
    if "initiatives" in data:
        add_heading_styled(doc, "Initiatives", level=1)
        headers = ["#", "Initiative", "Priority", "Timeline", "Budget"]
        add_data_table(doc, headers, data["initiatives"])

    # ─── Resource Loading ───
    if "resource_loading" in data:
        add_heading_styled(doc, "Resource Loading", level=1)
        add_resource_loading(doc, data["resource_loading"])

    # ─── 4D Mix ───
    if "four_d_mix" in data:
        add_heading_styled(doc, "4D Mix Analysis", level=1)
        for person in data["four_d_mix"]:
            add_four_d_table(doc, person["name"], person["current"],
                              person["ideal"], person["quarterly"])

    # ─── Execution Roadmap ───
    if "roadmap" in data:
        add_heading_styled(doc, "Execution Roadmap", level=1)
        for quarter in data["roadmap"]:
            add_heading_styled(doc, quarter["name"], level=2)
            add_bullet_list(doc, quarter["items"])
            if "revenue_target" in quarter:
                add_styled_paragraph(
                    doc, f"Revenue target: {quarter['revenue_target']}",
                    size=10, bold=True, color=MED_BLUE
                )

    # ─── Vision Alignment Score ───
    if "vision_alignment" in data:
        add_heading_styled(doc, "Vision Alignment Score", level=1)
        va = data["vision_alignment"]
        add_data_table(doc, ["Criteria", "Score"], va.get("scores", []))

        if "narrative" in va:
            add_insight_box(doc, va["narrative"])

    # ─── Appendix ───
    add_heading_styled(doc, "Appendix", level=1)
    add_styled_paragraph(
        doc,
        "Prepared using BSC Strategic Planning methodology with "
        "AI-human collaborative analysis.",
        size=9, italic=True, color=GRAY
    )

    doc.save(output_path)
    return output_path


# ── DATA TEMPLATE ──────────────────────────────────────────────────
# AI: Populate this dict with actual content from bsc-data/*.md files
# then call export_bsc_plan(DATA, output_path)

DATA_TEMPLATE = {
    "company_name":  "{COMPANY NAME}",
    "doc_title":     "BSC Strategic Plan",
    "subtitle":      "BSC Strategic Plan",
    "period":        "2025-2026 (1 year)",
    "date":          "13/03/2026",
    "prepared_by":   "BSC Strategic Planning Consultant",
    "prepared_for":  "Ban Lãnh đạo",

    "executive_summary": "...",

    "dashboard": [
        ["Revenue Target",     "2.4B VND"],
        ["Core Strategies",    "3"],
        ["Strategic Objectives", "17"],
    ],
    "dashboard_intro": "Tóm lược các chỉ số chiến lược chính của kế hoạch.",

    "company_profile": [
        "Company: ...",
        "Industry: ...",
    ],

    "vision":   "...",
    "mission":  "...",
    "vision_milestones": ["Y1: ...", "Y3: ...", "Y5+: ..."],

    "products": [
        {"name": "Product 1", "description": "...", "details": "..."},
    ],

    "reality_check": {
        "intro":               "Bối cảnh thị trường và những khoảng trống nhận thức.",
        "market_context":      "...",
        "competitive_position": "...",
        "critical_risks": [
            ["Risk 1", "High", "Mitigation..."],
        ],
        "perception_gaps": [
            ["Dim", "User said...", "Research found...", "Gap"],
        ],
        "key_insight": "...",
    },

    "swot": {
        "strengths":     ["S1...", "S2..."],
        "weaknesses":    ["W1...", "W2..."],
        "opportunities": ["O1...", "O2..."],
        "threats":       ["T1...", "T2..."],
        "insight":       "...",
    },

    "strategies_intro": "Ba chiến lược cốt lõi được chọn dựa trên phân tích TOWS.",
    "strategies": [
        {"name": "CS1: ...", "type": "Direct", "tows": "S-O",
         "rationale": "..."},
    ],

    "objectives": [
        {
            "name":      "Financial",
            "items": [["F1", "Revenue 2.4B", "Cumulative", "300M", "900M", "2.4B"]],
            "causality": "...",
        },
    ],

    "causal_chains": [
        {"id": 1, "foundation": "AI+SOPs", "process": "Cycle time",
         "customer": "Conv rate", "financial": "Revenue"},
    ],
    "chain_narrative": "...",

    "initiatives": [
        ["I1", "SOP to AI Engine", "P0", "Q1", "Internal"],
    ],

    "resource_loading": [
        {
            "name": "CEO",
            "role": "CEO",
            "allocations": [
                {"activity": "Sales",     "percent": 25},
                {"activity": "Strategic", "percent": 15},
            ],
            "total":  100,
            "status": "Balanced",
        },
    ],

    "four_d_mix": [
        {
            "name":    "CEO",
            "current": {"do": 45, "decide": 30, "delegate": 15, "design": 10},
            "ideal":   {"do": 10, "decide": 25, "delegate": 25, "design": 40},
            "quarterly": [
                {"quarter": "Q1", "do": 35, "decide": 28,
                 "delegate": 20, "design": 17},
            ],
        },
    ],

    "roadmap": [
        {
            "name":           "Q1 (Apr-Jun 2026)",
            "items":          ["Item 1", "Item 2"],
            "revenue_target": "300M cumulative",
        },
    ],

    "vision_alignment": {
        "scores": [
            ["Objectives traced to Vision", "17/17 = 100%"],
            ["Overall Alignment",           "STRONG (100%)"],
        ],
        "narrative": "...",
    },
}


# ── Entry Point ────────────────────────────────────────────────────

if __name__ == "__main__":
    output = "./bsc-data/BSC-Strategic-Plan-DEMO.docx"
    os.makedirs("./bsc-data", exist_ok=True)
    export_bsc_plan(DATA_TEMPLATE, output)
    print(f"Exported: {output}")
